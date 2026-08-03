"""Policy documents, loaded as LangChain Documents.

Implemented as a `BaseLoader` rather than using a stock loader because the
built-in ones carry only a path and a page number, and citations here need more
than that: which category a document belongs to for the library filters, and
which organisation issued it so a reader can see where an answer came from.
"""
from pathlib import Path
from typing import Iterator, List

from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document

from rag_engine.config import POLICY_DOCUMENTS_DIR

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# The bundled corpus is a set of publicly published policies from several real
# organisations, used as sample HR content. Attributing each document to its
# issuer keeps citations honest — a reader can see the answer came from, say,
# Bandhan's manual rather than from the demo employer's own handbook.
_COMPANY_KEYWORDS = (
    ("orbis", "Orbis"),
    ("bandhan", "Bandhan Financial Services"),
    ("ppl_", "Piramal Pharma"),
    ("pel ", "Piramal Enterprises"),
    ("pel-", "Piramal Enterprises"),
    ("pchfl", "Piramal Capital & Housing Finance"),
    ("smp", "Piramal Capital & Housing Finance"),
    ("posh", "Piramal Capital & Housing Finance"),
    ("remuneration-policy-new", "Piramal Housing Finance"),
)

UNKNOWN_COMPANY = "Unattributed"


def infer_company(file_name: str) -> str:
    """Best-effort issuer for a policy file, shown alongside citations.
    Uploaded files that match nothing are simply left unattributed."""
    name = file_name.lower()
    for keyword, company in _COMPANY_KEYWORDS:
        if keyword in name:
            return company
    return UNKNOWN_COMPANY


def infer_category(file_name: str) -> str:
    """Classify a policy file into the categories the UI filters by:
    Leave, Conduct, Benefits, Privacy, Work, General."""
    name = file_name.lower()

    if any(k in name for k in ("leave", "maternity", "paternity")):
        return "Leave"
    if any(k in name for k in ("posh", "harassment", "conduct", "ethics", "disciplin", "whistle")):
        return "Conduct"
    if any(k in name for k in ("benefit", "health", "wellness", "wellbeing",
                               "remuneration", "salary", "travel", "expense",
                               "reimburse", "learning", "development",
                               "certification", "insurance", "safety")):
        return "Benefits"
    if any(k in name for k in ("privacy", "confidential", "data", "diba")):
        return "Privacy"
    if any(k in name for k in ("work", "hybrid", "remote", "wfh", "lifecycle",
                               "probation", "notice", "exit", "resignation",
                               "onboarding", "attendance")):
        return "Work"
    return "General"


class PolicyDocumentLoader(BaseLoader):
    """Load one policy file as a Document per page.

    Splitting by page rather than by file is what allows a citation to name a
    page; a whole-document Document would lose that.
    """

    def __init__(self, file_path: Path | str):
        self.path = Path(file_path)

    def _metadata(self, page: int) -> dict:
        return {
            "source": self.path.name,
            "page": page,
            "category": infer_category(self.path.name),
            "company": infer_company(self.path.name),
        }

    def lazy_load(self) -> Iterator[Document]:
        suffix = self.path.suffix.lower()
        if suffix == ".pdf":
            yield from self._load_pdf()
        elif suffix == ".docx":
            yield from self._load_docx()
        elif suffix == ".txt":
            yield from self._load_txt()

    def _load_pdf(self) -> Iterator[Document]:
        from pypdf import PdfReader

        for number, page in enumerate(PdfReader(str(self.path)).pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:  # scanned or empty pages carry nothing to retrieve
                yield Document(page_content=text, metadata=self._metadata(number))

    def _load_docx(self) -> Iterator[Document]:
        from docx import Document as DocxDocument

        paragraphs = [p.text.strip() for p in DocxDocument(str(self.path)).paragraphs]
        text = "\n".join(p for p in paragraphs if p)
        if text:
            yield Document(page_content=text, metadata=self._metadata(1))

    def _load_txt(self) -> Iterator[Document]:
        text = self.path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            yield Document(page_content=text, metadata=self._metadata(1))


def load_document(file_path: Path | str) -> List[Document]:
    """Every page of a single policy file."""
    return list(PolicyDocumentLoader(file_path).lazy_load())


def load_policy_documents(policy_dir: Path = POLICY_DOCUMENTS_DIR) -> List[Document]:
    """Every page of every policy file in the folder."""
    if not policy_dir.exists():
        raise FileNotFoundError(f"Policy documents folder not found: {policy_dir}")

    documents: List[Document] = []
    for path in sorted(policy_dir.iterdir()):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            documents.extend(load_document(path))
    return documents
